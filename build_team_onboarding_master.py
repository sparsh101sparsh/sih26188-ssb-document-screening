import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import subprocess

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def build_docx(md_path, docx_out_paths):
    print("Building publication-grade Word document (.docx)...")
    doc = docx.Document()
    
    # 0.75 in margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Header & Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("SIH26188: AI-Based Fake Identity Screening — 30-Phase Team Onboarding Guide | MHA / SSB")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = RGBColor(148, 163, 184)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("MINISTRY OF HOME AFFAIRS • SASHASTRA SEEMA BAL (SSB) — SMART INDIA HACKATHON 2026")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(8)
        frun.font.color.rgb = RGBColor(148, 163, 184)

    def add_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 37, 82)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 118, 110) # Teal
        return p

    def add_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(14 if level==1 else (10 if level==2 else 6))
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(15, 37, 82)
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(15, 118, 110)
        else:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(30, 41, 59)
        return p

    def add_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix + " ")
            r_bold.font.name = 'Calibri'
            r_bold.font.bold = True
            r_bold.font.size = Pt(10)
            r_bold.font.color.rgb = RGBColor(15, 23, 42)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(30, 41, 59)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix + " ")
            r_bold.font.name = 'Calibri'
            r_bold.font.bold = True
            r_bold.font.size = Pt(9.5)
            r_bold.font.color.rgb = RGBColor(15, 23, 42)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_callout(text, title="OPERATIONAL CONTEXT"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(7.0)
        set_cell_background(cell, "F0F9FF") # Light blue
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="none"/>'
            f'<w:left w:val="single" w:sz="24" w:color="0284C7"/>'
            f'<w:bottom w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if title:
            r_title = p.add_run(f"📌 {title}\n")
            r_title.font.name = 'Calibri'
            r_title.font.bold = True
            r_title.font.size = Pt(9.5)
            r_title.font.color.rgb = RGBColor(2, 132, 199)
        
        r_body = p.add_run(text)
        r_body.font.name = 'Calibri'
        r_body.font.size = Pt(9.5)
        r_body.font.color.rgb = RGBColor(30, 41, 59)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(7.0)
        set_cell_background(cell, "F8FAFC") # Slate light
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:color="CBD5E1"/>'
            f'<w:left w:val="single" w:sz="12" w:color="94A3B8"/>'
            f'<w:bottom w:val="single" w:sz="4" w:color="CBD5E1"/>'
            f'<w:right w:val="single" w:sz="4" w:color="CBD5E1"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(code_text)
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(15, 23, 42)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_table(headers, rows, col_widths=None):
        tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        set_table_borders(tbl, color="CBD5E1", sz="4", val="single")
        
        hdr_row = tbl.rows[0]
        for i, header_text in enumerate(headers):
            cell = hdr_row.cells[i]
            if col_widths and i < len(col_widths):
                cell.width = Inches(col_widths[i])
            set_cell_background(cell, "1E3A8A")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(header_text)
            r.font.name = 'Calibri'
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(255, 255, 255)
            
        for row_idx, row_data in enumerate(rows):
            row = tbl.rows[row_idx + 1]
            bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            for col_idx, cell_value in enumerate(row_data):
                cell = row.cells[col_idx]
                if col_widths and col_idx < len(col_widths):
                    cell.width = Inches(col_widths[col_idx])
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(cell_value)
                r.font.name = 'Calibri'
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(30, 41, 59)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # Read Markdown and render
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_buffer = []
    in_table = False
    table_headers = []
    table_rows = []

    for line in lines:
        raw_line = line.rstrip()
        
        # Code block handling
        if raw_line.startswith("```"):
            if in_code_block:
                add_code_block("\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_buffer.append(raw_line)
            continue

        # Table handling
        if raw_line.startswith("|") and raw_line.endswith("|"):
            cells = [c.strip() for c in raw_line.strip("|").split("|")]
            if all(set(c).issubset({'-', ':', ' '}) for c in cells):
                continue # delimiter line
            if not in_table:
                in_table = True
                table_headers = cells
                table_rows = []
            else:
                table_rows.append(cells)
            continue
        else:
            if in_table:
                if table_headers and table_rows:
                    add_table(table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []

        if not raw_line:
            continue

        # Headings
        if raw_line.startswith("# "):
            add_title(raw_line[2:].strip())
        elif raw_line.startswith("## "):
            add_heading(raw_line[3:].strip(), level=1)
        elif raw_line.startswith("### "):
            add_heading(raw_line[4:].strip(), level=2)
        elif raw_line.startswith("#### "):
            add_heading(raw_line[5:].strip(), level=3)
        elif raw_line.startswith("> "):
            add_callout(raw_line[2:].strip())
        elif raw_line.startswith("- ") or raw_line.startswith("* "):
            parts = raw_line[2:].split("**")
            if len(parts) >= 3:
                bold_pfx = parts[1]
                rest_txt = "".join(parts[2:])
                add_bullet(rest_txt.lstrip(": ").strip(), bold_prefix=bold_pfx)
            else:
                add_bullet(raw_line[2:].strip())
        elif raw_line.startswith("1. ") or raw_line.startswith("2. ") or raw_line.startswith("3. ") or raw_line.startswith("4. ") or raw_line.startswith("5. "):
            pfx = raw_line[:3]
            body = raw_line[3:]
            parts = body.split("**")
            if len(parts) >= 3:
                bold_pfx = pfx + " " + parts[1]
                rest_txt = "".join(parts[2:])
                add_bullet(rest_txt.lstrip(": ").strip(), bold_prefix=bold_pfx)
            else:
                add_bullet(body.strip(), bold_prefix=pfx)
        else:
            parts = raw_line.split("**")
            if len(parts) >= 3 and raw_line.startswith("**"):
                bold_pfx = parts[1]
                rest_txt = "".join(parts[2:])
                add_p(rest_txt.lstrip(": ").strip(), bold_prefix=bold_pfx)
            else:
                add_p(raw_line)

    if in_table and table_headers and table_rows:
        add_table(table_headers, table_rows)

    for out_path in docx_out_paths:
        doc.save(out_path)
        print(f"Saved DOCX: {out_path} ({os.path.getsize(out_path)} bytes)")

def build_pdf_with_typst(md_path, pdf_out_paths):
    print("Building high-clarity PDF via Typst...")
    typ_file = "/Users/iamsparsh00321/Documents/antigravity/vibrant-rutherford/temp_styled_onboarding.typ"
    
    # Convert MD to Typst
    subprocess.run(["pandoc", md_path, "-o", typ_file], check=True)
    
    with open(typ_file, "r", encoding="utf-8") as f:
        typ_content = f.read()

    # Prepend elegant Typst layout styles
    custom_header = """
#set page(
  paper: "a4",
  margin: (top: 2cm, bottom: 2cm, left: 2cm, right: 2cm),
  header: align(right)[
    #text(7.5pt, fill: rgb("#64748b"))[
      *SIH26188 — Ministry of Home Affairs / SSB AI Document Screening*
    ]
  ],
  footer: [
    #line(length: 100%, stroke: 0.5pt + rgb("#cbd5e1"))
    #grid(
      columns: (1fr, 1fr),
      align(left)[#text(7.5pt, fill: rgb("#94a3b8"))[Sashastra Seema Bal (SSB) • Police II Division]],
      align(right)[#text(7.5pt, fill: rgb("#94a3b8"))[Page #context counter(page).display()]]
    )
  ]
)

#set text(
  font: ("Arial", "Helvetica"),
  size: 9.5pt,
  fill: rgb("#0f172a"),
  lang: "en"
)

#set par(justify: true, leading: 0.6em)

#show heading.where(level: 1): it => block(
  above: 1.4em, below: 0.7em,
  text(weight: "bold", size: 14.5pt, fill: rgb("#1e3a8a"))[#it.body]
)

#show heading.where(level: 2): it => block(
  above: 1.2em, below: 0.5em,
  text(weight: "bold", size: 12pt, fill: rgb("#0f766e"))[#it.body]
)

#show heading.where(level: 3): it => block(
  above: 0.9em, below: 0.3em,
  text(weight: "bold", size: 10.5pt, fill: rgb("#1e293b"))[#it.body]
)

#show raw.where(block: true): it => block(
  fill: rgb("#f8fafc"),
  inset: 7pt,
  radius: 3pt,
  stroke: 0.5pt + rgb("#cbd5e1"),
  width: 100%,
  text(size: 7pt, fill: rgb("#0f172a"), font: "Courier New")[#it]
)

#show table: set text(size: 8pt)
#set table(stroke: 0.5pt + rgb("#cbd5e1"), inset: 4.5pt)

#show quote: it => block(
  fill: rgb("#f0f9ff"),
  inset: (x: 10pt, y: 7pt),
  radius: 3pt,
  stroke: (left: 3pt + rgb("#0284c7")),
  width: 100%,
  text(size: 9pt, fill: rgb("#0f172a"))[#it.body]
)
"""
    with open(typ_file, "w", encoding="utf-8") as f:
        f.write(custom_header + "\n" + typ_content)

    temp_pdf = "/Users/iamsparsh00321/Documents/antigravity/vibrant-rutherford/temp_styled_onboarding.pdf"
    subprocess.run(["typst", "compile", typ_file, temp_pdf], check=True)

    for p in pdf_out_paths:
        subprocess.run(["cp", temp_pdf, p], check=True)
        print(f"Saved PDF: {p} ({os.path.getsize(p)} bytes)")

    if os.path.exists(typ_file):
        os.remove(typ_file)
    if os.path.exists(temp_pdf):
        os.remove(temp_pdf)

if __name__ == "__main__":
    md_file = "/Users/iamsparsh00321/Documents/antigravity/vibrant-rutherford/SIH26188_TEAM_ONBOARDING_30_PHASES.md"
    docx_paths = [
        "/Users/iamsparsh00321/Documents/antigravity/vibrant-rutherford/SIH26188_Team_Onboarding_30_Phases.docx",
        "/Users/iamsparsh00321/Documents/antigravity/vibrant-rutherford/sih26188_project/docs/SIH26188_Team_Onboarding_30_Phases.docx"
    ]
    pdf_paths = [
        "/Users/iamsparsh00321/Documents/antigravity/vibrant-rutherford/SIH26188_Team_Onboarding_30_Phases.pdf",
        "/Users/iamsparsh00321/Documents/antigravity/vibrant-rutherford/sih26188_project/docs/SIH26188_Team_Onboarding_30_Phases.pdf"
    ]
    
    build_docx(md_file, docx_paths)
    build_pdf_with_typst(md_file, pdf_paths)
    print("ALL BUILDS COMPLETED SUCCESSFULLY!")
